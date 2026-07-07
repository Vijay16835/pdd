import os
import traceback
from typing import Optional

# ---------------------------------------------------------------------------
# Supabase singleton — created once at module load, reused on every request.
# ---------------------------------------------------------------------------
_supabase_client = None

def get_supabase():
    """Return the shared Supabase client, initialising it on first call."""
    global _supabase_client
    if _supabase_client is None:
        from supabase import create_client
        from app.core.config import settings
        _supabase_client = create_client(settings.SUPABASE_URL, settings.SUPABASE_KEY)
    return _supabase_client


class TextExtractionError(Exception):
    """Custom exception raised when text extraction fails."""
    pass


def extract_text_from_pdf(file_path: str) -> str:
    """Extract text from PDF files using pdfplumber."""
    try:
        import pdfplumber
        text = ""
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n\n"
        return text.strip()
    except Exception as e:
        print(f"PDF extraction error: {e}")
        # Fallback to PyPDF2
        try:
            import PyPDF2
            text = ""
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n\n"
            return text.strip()
        except Exception as e2:
            print(f"PyPDF2 fallback failed: {e2}")
            raise TextExtractionError("PDF text extraction failed") from e2


def extract_text_from_doc(file_path: str) -> str:
    """Extract text from legacy DOC files using legacy-doc."""
    try:
        from legacy_doc import extract_text as legacy_extract
        with open(file_path, "rb") as f:
            result = legacy_extract(f.read())
            if hasattr(result, 'text'):
                text = result.text
            else:
                text = str(result)
        if not text.strip():
            raise TextExtractionError("Unable to extract readable text from document.")
        return text.strip()
    except TextExtractionError:
        raise
    except Exception as e:
        print(f"DOC extraction error: {e}")
        traceback.print_exc()
        raise TextExtractionError("Unsupported file structure") from e


def extract_text_from_docx(file_path: str) -> str:
    """Extract text from DOCX files."""
    try:
        import docx
        doc = docx.Document(file_path)
        text = ""
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text += paragraph.text + "\n"
        
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text += row_text + "\n"
        
        if not text.strip():
            raise TextExtractionError("Unable to extract readable text from document.")
            
        return text.strip()
    except TextExtractionError:
        raise
    except Exception as e:
        print(f"DOCX extraction error: {e}")
        traceback.print_exc()
        raise TextExtractionError("Unsupported file structure") from e


def extract_text_from_txt(file_path: str) -> str:
    """Extract text from plain text files."""
    try:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            text = f.read().strip()
        if not text:
            raise TextExtractionError("Unable to extract readable text from document.")
        return text
    except TextExtractionError:
        raise
    except Exception as e:
        print(f"TXT extraction error: {e}")
        raise TextExtractionError("Unsupported file structure") from e


def preprocess_image(file_path: str) -> str:
    """
    Preprocess the image before running OCR.
    Steps:
      1. Open with Pillow to handle EXIF orientation and format conversions.
      2. Convert to RGB mode and then to OpenCV BGR format.
      3. Resize if image is too large (max dimension > 2000px)
      4. Convert to grayscale
      5. Increase contrast using CLAHE
      6. Denoise using fastNlMeansDenoising
    Saves to a temporary file and returns its path.
    """
    import cv2
    import numpy as np
    import tempfile
    from PIL import Image, ImageOps
    
    # Read image using PIL to handle EXIF orientation and format conversions
    try:
        with Image.open(file_path) as pil_img:
            # Correct orientation based on EXIF data
            pil_img = ImageOps.exif_transpose(pil_img)
            # Convert to RGB mode (removes alpha channel, handles palettes)
            pil_img = pil_img.convert("RGB")
            # Convert to OpenCV BGR format
            img = cv2.cvtColor(np.array(pil_img), cv2.COLOR_RGB2BGR)
    except Exception as e:
        print(f"PIL image loading failed: {e}. Falling back to cv2.imread.")
        # Fallback to direct cv2 read
        img = cv2.imread(file_path)
        
    if img is None:
        raise TextExtractionError("Unable to read image file.")
        
    # 1. Resize large images (keeps aspect ratio)
    h, w = img.shape[:2]
    max_dim = 2000
    if max(h, w) > max_dim:
        scale = max_dim / max(h, w)
        img = cv2.resize(img, (int(w * scale), int(h * scale)), interpolation=cv2.INTER_AREA)
        
    # 2. Convert to grayscale
    gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
    
    # 3. Increase contrast using CLAHE (Contrast Limited Adaptive Histogram Equalization)
    clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8, 8))
    contrast = clahe.apply(gray)
    
    # 4. Denoise image (milder strength to prevent smudging thin fonts)
    denoised = cv2.fastNlMeansDenoising(contrast, None, h=3, templateWindowSize=7, searchWindowSize=21)
    
    # Save to a temporary file
    with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp:
        temp_path = tmp.name
        
    cv2.imwrite(temp_path, denoised)
    return temp_path


def extract_text_from_image(file_path: str) -> str:
    """Extract text from images using OCR (EasyOCR or pytesseract) with preprocessing and confidence validation."""
    print("[IMAGE] Image received")
    print("[OCR] OCR started")
    
    preprocessed_path = None
    ocr_target_path = file_path
    try:
        preprocessed_path = preprocess_image(file_path)
        ocr_target_path = preprocessed_path
    except Exception as pe:
        print(f"Image preprocessing warning: {pe}")
        
    text = ""
    easyocr_success = False
    
    # Try EasyOCR first
    try:
        import easyocr
        reader = easyocr.Reader(['en'], gpu=False)
        # detail=1 returns bounding box, text, and confidence score
        results = reader.readtext(ocr_target_path, detail=1)
        
        if results:
            confidences = [res[2] for res in results]
            texts = [res[1] for res in results]
            
            # Check average confidence
            avg_confidence = sum(confidences) / len(confidences) if confidences else 0.0
            LOW_CONFIDENCE_THRESHOLD = 0.40
            if avg_confidence >= LOW_CONFIDENCE_THRESHOLD:
                text = "\n".join(texts)
                if text.strip():
                    easyocr_success = True
            else:
                print(f"EasyOCR low confidence: {avg_confidence:.2f} < {LOW_CONFIDENCE_THRESHOLD}")
    except Exception as e:
        print(f"EasyOCR failed: {e}")
    
    # Fallback to pytesseract if EasyOCR failed, had low confidence, or returned empty text
    if not easyocr_success:
        print("[OCR] EasyOCR unavailable or low confidence. Trying pytesseract fallback...")
        try:
            from app.core.config import settings
            import pytesseract
            from PIL import Image

            # Set tesseract path if it exists
            if os.path.exists(settings.TESSERACT_CMD):
                pytesseract.pytesseract.tesseract_cmd = settings.TESSERACT_CMD

            img = Image.open(ocr_target_path)

            # Confidence check
            t_conf_ok = True
            try:
                data = pytesseract.image_to_data(img, output_type=pytesseract.Output.DICT)
                confidences = [float(c) for c in data['conf'] if c != '-1' and c != -1]
                if confidences:
                    avg_confidence = sum(confidences) / (len(confidences) * 100.0)
                    LOW_CONFIDENCE_THRESHOLD = 0.40
                    if avg_confidence < LOW_CONFIDENCE_THRESHOLD:
                        print(f"Pytesseract low confidence: {avg_confidence:.2f} < {LOW_CONFIDENCE_THRESHOLD}")
                        t_conf_ok = False
            except Exception as t_conf_err:
                print(f"Pytesseract confidence check warning: {t_conf_err}")

            t_text = pytesseract.image_to_string(img)
            if t_text.strip() and t_conf_ok:
                text = t_text
            elif t_text.strip() and not easyocr_success:
                # Pytesseract has low confidence but EasyOCR gave nothing — use it anyway
                text = t_text
        except Exception as e:
            print(f"Pytesseract failed: {e}")
            # Only raise if we truly have nothing from either engine
            if not text.strip():
                raise TextExtractionError("OCR extraction failed — no text could be extracted from the image.") from e
    
    # Clean up preprocessed file
    if preprocessed_path and os.path.exists(preprocessed_path):
        try:
            os.remove(preprocessed_path)
        except Exception as remove_err:
            print(f"Failed to remove temporary preprocessed image: {remove_err}")
            
    stripped_text = text.strip()
    if not stripped_text:
        raise TextExtractionError("No readable text found in image.")
        
    print("[OCR] OCR completed")
    print(f"[OCR] Extracted text length = {len(stripped_text)}")
    return stripped_text


def ocr_pdf(file_path: str) -> str:
    """Perform OCR on a PDF by rendering pages as images and running OCR on them."""
    import pypdfium2 as pdfium
    import tempfile
    
    print("[OCR] OCR started")
    text = ""
    pdf = None
    try:
        pdf = pdfium.PdfDocument(file_path)
        for i in range(len(pdf)):
            page = pdf.get_page(i)
            pil_image = page.render(scale=2).to_pil()
            
            # Save page image to temporary file
            with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp:
                temp_img_path = tmp.name
            try:
                pil_image.save(temp_img_path)
                page_text = extract_text_from_image(temp_img_path)
                if page_text:
                    text += page_text + "\n\n"
            finally:
                if os.path.exists(temp_img_path):
                    os.remove(temp_img_path)
    except Exception as e:
        print(f"ocr_pdf failed: {e}")
        raise TextExtractionError("OCR extraction failed") from e
    finally:
        if pdf:
            pdf.close()
            
    print("[OCR] OCR completed")
    if not text.strip():
        raise TextExtractionError("OCR extraction failed")
    return text.strip()


def extract_text(file_path: str, file_type: str) -> str:
    """Main dispatcher — extract text based on file type."""
    file_type = file_type.lower()
    print("[EXTRACT] Starting extraction")
    
    text = ""
    try:
        if file_type in ['pdf', 'application/pdf']:
            text = extract_text_from_pdf(file_path)
        elif file_type in ['doc', 'docx', 'application/msword', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document']:
            if file_type in ['doc', 'application/msword'] or file_path.endswith('.doc'):
                text = extract_text_from_doc(file_path)
            else:
                text = extract_text_from_docx(file_path)
        elif file_type in ['txt', 'text/plain']:
            text = extract_text_from_txt(file_path)
        elif file_type in ['jpg', 'jpeg', 'png', 'image/jpeg', 'image/png', 'image/jpg']:
            text = extract_text_from_image(file_path)
        else:
            raise TextExtractionError("Unsupported file structure")
    except TextExtractionError:
        raise
    except Exception as e:
        print(f"Extraction failed: {e}")
        traceback.print_exc()
        if file_type in ['pdf', 'application/pdf']:
            raise TextExtractionError("PDF text extraction failed") from e
        elif file_type in ['doc', 'docx', 'application/msword', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document']:
            raise TextExtractionError("Unsupported file structure") from e
        elif file_type in ['jpg', 'jpeg', 'png', 'image/jpeg', 'image/png', 'image/jpg']:
            raise TextExtractionError("OCR extraction failed") from e
        else:
            raise TextExtractionError("Unsupported file structure") from e

    # OCR Fallback if text is empty
    if not text.strip():
        print("[EXTRACT] Extracted text empty. Starting OCR fallback.")
        if file_type in ['pdf', 'application/pdf']:
            try:
                text = ocr_pdf(file_path)
            except Exception as ocr_err:
                raise TextExtractionError("OCR extraction failed") from ocr_err
        elif file_type in ['jpg', 'jpeg', 'png', 'image/jpeg', 'image/png', 'image/jpg']:
            try:
                text = extract_text_from_image(file_path)
            except Exception as ocr_err:
                raise TextExtractionError("OCR extraction failed") from ocr_err
        else:
            raise TextExtractionError("Unable to extract readable text from document.")
            
    if not text.strip():
        if file_type in ['pdf', 'application/pdf'] or file_type in ['jpg', 'jpeg', 'png', 'image/jpeg', 'image/png', 'image/jpg']:
            raise TextExtractionError("OCR extraction failed")
        else:
            raise TextExtractionError("Unable to extract readable text from document.")
            
    print(f"[EXTRACT] Text length: {len(text)}")
    return text.strip()


def get_file_extension(filename: str) -> str:
    """Get clean file extension from filename."""
    _, ext = os.path.splitext(filename)
    return ext.lstrip('.').lower()


SUPPORTED_EXTENSIONS = {'pdf', 'doc', 'docx', 'txt', 'jpg', 'jpeg', 'png'}
MAX_FILE_SIZE_MB = 20

def validate_file(filename: str, file_size: int) -> tuple:
    """Validate file type and size. Returns (is_valid, error_message)."""
    ext = get_file_extension(filename)
    if ext not in SUPPORTED_EXTENSIONS:
        return False, f"Unsupported file type: .{ext}. Supported: {', '.join(SUPPORTED_EXTENSIONS)}"
    
    size_mb = file_size / (1024 * 1024)
    if size_mb > MAX_FILE_SIZE_MB:
        return False, f"File too large: {size_mb:.1f}MB. Maximum: {MAX_FILE_SIZE_MB}MB"
    
    return True, ""


def get_user_storage_usage_mb(user_id: str) -> float:
    """Calculate the total storage used by a user in MB.
    Includes:
    - Uploaded files (from PostgreSQL/Firestore, checking local disk or Supabase Storage size)
    - Generated reports (from local disk reports directory matching user's document IDs)
    """
    total_size_bytes = 0
    doc_ids_and_sizes = {}  # doc_id -> size_in_bytes

    # 1. Fetch documents from PostgreSQL
    try:
        from app.services.firebase_service import firebase_service
        conn = firebase_service._get_pg_conn()
        if conn:
            try:
                cur = conn.cursor()
                cur.execute("SELECT id, size_in_mb FROM documents WHERE user_id = %s", (user_id,))
                rows = cur.fetchall()
                for row in rows:
                    doc_id = row[0]
                    size_mb = row[1] or 0.0
                    doc_ids_and_sizes[doc_id] = int(size_mb * 1024 * 1024)
                cur.close()
                conn.close()
            except Exception as e:
                print(f"Error querying PostgreSQL documents: {e}")
                if conn:
                    conn.close()
    except Exception as e:
        print(f"PostgreSQL connection error: {e}")

    # 2. Fetch documents from Firestore
    try:
        from app.services.firebase_service import firebase_service
        if firebase_service.db:
            docs = firebase_service.db.collection("documents").where("user_id", "==", user_id).stream()
            for doc in docs:
                doc_id = doc.id
                data = doc.to_dict()
                size_mb = data.get("size_in_mb", 0.0) or 0.0
                if doc_id not in doc_ids_and_sizes:
                    doc_ids_and_sizes[doc_id] = int(size_mb * 1024 * 1024)
    except Exception as e:
        print(f"Error querying Firestore documents: {e}")

    # 3. Fetch documents from Supabase Storage
    try:
        supabase = get_supabase()
        files = supabase.storage.from_("legal-documents").list(f"users/{user_id}/documents")
        if files:
            for f in files:
                name = f.get("name")
                if name:
                    doc_id, ext = os.path.splitext(name)
                    metadata = f.get('metadata')
                    if metadata:
                        size_bytes = metadata.get('size', 0)
                        doc_ids_and_sizes[doc_id] = size_bytes
    except Exception as e:
        print(f"Error listing Supabase Storage files: {e}")

    # 4. Check actual file sizes on local disk
    from app.core.config import settings
    upload_dir = settings.UPLOAD_DIR
    local_doc_sizes = {}
    if os.path.exists(upload_dir):
        try:
            for filename in os.listdir(upload_dir):
                filepath = os.path.join(upload_dir, filename)
                if os.path.isfile(filepath):
                    doc_id, ext = os.path.splitext(filename)
                    if doc_id in doc_ids_and_sizes:
                        local_doc_sizes[doc_id] = os.path.getsize(filepath)
        except Exception as e:
            print(f"Error scanning local upload directory: {e}")

    # For each found document ID, use local size if it exists, else the db/remote size
    for doc_id, remote_size in doc_ids_and_sizes.items():
        if doc_id in local_doc_sizes:
            total_size_bytes += local_doc_sizes[doc_id]
        else:
            total_size_bytes += remote_size

    # 5. Check generated reports
    # Reports are under settings.UPLOAD_DIR/reports/LexGuard_Analysis_{doc_id}.{ext}
    reports_dir = os.path.join(upload_dir, "reports")
    if os.path.exists(reports_dir):
        try:
            for filename in os.listdir(reports_dir):
                filepath = os.path.join(reports_dir, filename)
                if os.path.isfile(filepath):
                    # Check if this report belongs to any of the user's document IDs
                    for doc_id in doc_ids_and_sizes.keys():
                        if f"LexGuard_Analysis_{doc_id}" in filename:
                            total_size_bytes += os.path.getsize(filepath)
                            break
        except Exception as e:
            print(f"Error scanning local reports directory: {e}")

    total_size_mb = total_size_bytes / (1024 * 1024)
    return round(total_size_mb, 2)
